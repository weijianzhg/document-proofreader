from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import io

class DocumentProcessor:
    def __init__(self, uploaded_file):
        """Initialize with uploaded file"""
        self.uploaded_file = uploaded_file
        self.document = Document(uploaded_file)
    
    def extract_paragraphs(self):
        """Extract all paragraphs from the document"""
        paragraphs = []
        
        for paragraph in self.document.paragraphs:
            # Get the full text of the paragraph
            text = paragraph.text
            paragraphs.append(text)
        
        return paragraphs
    
    def create_final_document(self, corrected_paragraphs):
        """Create a new document with corrected paragraphs"""
        # Create a new document based on the original
        new_doc = Document()
        
        # Copy styles from original document if possible
        try:
            # Copy document properties
            new_doc.core_properties.title = self.document.core_properties.title
            new_doc.core_properties.author = self.document.core_properties.author
        except:
            pass
        
        # Add corrected paragraphs to new document
        for i, (original_para, corrected_text) in enumerate(zip(self.document.paragraphs, corrected_paragraphs)):
            # Create new paragraph in new document
            new_para = new_doc.add_paragraph()
            
            # Try to preserve original formatting
            try:
                # Copy paragraph style if it exists
                if original_para.style:
                    new_para.style = original_para.style
            except:
                pass
            
            # Add the corrected text
            run = new_para.add_run(corrected_text)
            
            # Try to preserve some basic formatting from the first run of original paragraph
            if original_para.runs:
                try:
                    first_run = original_para.runs[0]
                    if first_run.bold:
                        run.bold = first_run.bold
                    if first_run.italic:
                        run.italic = first_run.italic
                    if first_run.underline:
                        run.underline = first_run.underline
                    if first_run.font.size:
                        run.font.size = first_run.font.size
                    if first_run.font.name:
                        run.font.name = first_run.font.name
                except:
                    pass
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        new_doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
    
    def create_tracked_changes_document(self, original_paragraphs, corrected_paragraphs, approved_changes):
        """Create a document with tracked changes highlighting"""
        new_doc = Document()
        
        # Copy document properties
        try:
            new_doc.core_properties.title = self.document.core_properties.title + " - Tracked Changes"
            new_doc.core_properties.author = self.document.core_properties.author
        except:
            pass
        
        for i, (original_text, corrected_text) in enumerate(zip(original_paragraphs, corrected_paragraphs)):
            new_para = new_doc.add_paragraph()
            
            # If changes were made and approved
            if (i in approved_changes and approved_changes[i] and 
                original_text.strip() != corrected_text.strip() and original_text.strip()):
                
                # Add corrected text with highlighting
                run = new_para.add_run(corrected_text)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                # Add original text as strikethrough comment
                if original_text.strip():
                    comment_para = new_doc.add_paragraph()
                    comment_run = comment_para.add_run(f"Original: {original_text}")
                    comment_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
                    try:
                        comment_run.font.strike = True
                    except:
                        pass
            else:
                # Use original text
                new_para.add_run(original_text)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        new_doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes.getvalue()
